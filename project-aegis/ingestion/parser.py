"""
ingestion/parser.py

Converts raw enterprise documents into a unified markdown string.
Supported formats: PDF, DOCX, Markdown, HTML.

All downstream components (chunker, metadata extractor) work from this
normalised markdown representation regardless of the original source format.
"""

from __future__ import annotations

import re
from pathlib import Path

from loguru import logger


# ── Public entry-point ──────────────────────────────────────────────────────

def parse_document(file_path: str | Path) -> str:
    """
    Parse any supported document into a markdown string.

    Args:
        file_path: Path to the document on disk.

    Returns:
        Full document as a clean markdown string.

    Raises:
        ValueError: If the file extension is not supported.
    """
    path = Path(file_path)
    ext  = path.suffix.lower()
    logger.info(f"Parsing {path.name!r}  (format={ext})")

    dispatch = {
        ".pdf":  _parse_pdf,
        ".docx": _parse_docx,
        ".md":   _parse_markdown,
        ".html": _parse_html,
        ".htm":  _parse_html,
    }

    if ext not in dispatch:
        raise ValueError(
            f"Unsupported format: {ext!r}. "
            f"Accepted: {', '.join(dispatch)}"
        )

    raw     = dispatch[ext](path)
    cleaned = _clean(raw)
    logger.success(f"Parsed {path.name!r}: {len(cleaned):,} chars")
    return cleaned


# ── Format parsers ──────────────────────────────────────────────────────────

def _parse_pdf(path: Path) -> str:
    """
    Primary: UnstructuredPDFLoader (element-mode) converts headings,
    narrative text, and tables into structured markdown.
    Fallback: pdfminer plain-text extraction.
    """
    try:
        from langchain_community.document_loaders import UnstructuredPDFLoader

        loader = UnstructuredPDFLoader(str(path), mode="elements")
        docs   = loader.load()

        parts: list[str] = []
        for doc in docs:
            cat  = doc.metadata.get("category", "NarrativeText")
            text = doc.page_content.strip()
            if not text:
                continue
            if cat in ("Title", "Header"):
                parts.append(f"## {text}")
            elif cat == "Table":
                parts.append(_raw_table_to_md(text))
            else:
                parts.append(text)
        return "\n\n".join(parts)

    except ImportError:
        logger.warning("unstructured not installed — falling back to pdfminer")
        return _pdfminer_fallback(path)


def _pdfminer_fallback(path: Path) -> str:
    from pdfminer.high_level import extract_text
    return extract_text(str(path))


def _parse_docx(path: Path) -> str:
    """
    Walks paragraphs and tables in document order.
    Heading styles are mapped to # / ## / ### markers.
    Tables are converted to pipe-delimited markdown with a separator row.
    """
    from docx import Document

    doc   = Document(str(path))
    parts: list[str] = []

    for para in doc.paragraphs:
        text  = para.text.strip()
        style = para.style.name or ""
        if not text:
            continue
        if "Heading 1" in style:
            parts.append(f"# {text}")
        elif "Heading 2" in style:
            parts.append(f"## {text}")
        elif "Heading 3" in style:
            parts.append(f"### {text}")
        else:
            parts.append(text)

    for table in doc.tables:
        parts.append(_docx_table_to_md(table))

    return "\n\n".join(parts)


def _parse_markdown(path: Path) -> str:
    """Markdown files are returned as-is."""
    return path.read_text(encoding="utf-8")


def _parse_html(path: Path) -> str:
    """
    Strips HTML, preserving heading structure and table data as markdown.
    """
    from bs4 import BeautifulSoup

    soup  = BeautifulSoup(path.read_text(encoding="utf-8"), "html.parser")
    parts: list[str] = []

    for tag in soup.find_all(["h1", "h2", "h3", "p", "li", "table"]):
        text = tag.get_text(separator=" ").strip()
        if not text:
            continue
        if tag.name == "h1":
            parts.append(f"# {text}")
        elif tag.name == "h2":
            parts.append(f"## {text}")
        elif tag.name == "h3":
            parts.append(f"### {text}")
        elif tag.name == "li":
            parts.append(f"- {text}")
        elif tag.name == "table":
            parts.append(_html_table_to_md(tag))
        else:
            parts.append(text)

    return "\n\n".join(parts)


# ── Table converters ────────────────────────────────────────────────────────

def _docx_table_to_md(table) -> str:
    rows: list[str] = []
    for i, row in enumerate(table.rows):
        cells = [c.text.strip() for c in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
    return "\n".join(rows)


def _html_table_to_md(tag) -> str:
    rows: list[str] = []
    for i, tr in enumerate(tag.find_all("tr")):
        cells = [td.get_text(strip=True) for td in tr.find_all(["th", "td"])]
        rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
    return "\n".join(rows)


def _raw_table_to_md(raw: str) -> str:
    """Best-effort conversion of unstructured table text."""
    lines = [l.strip() for l in raw.splitlines() if l.strip()]
    if not lines:
        return raw
    md_rows = [f"| {l} |" for l in lines]
    sep     = "| --- |"
    return md_rows[0] + "\n" + sep + "\n" + "\n".join(md_rows[1:])


# ── Cleaning ────────────────────────────────────────────────────────────────

def _clean(text: str) -> str:
    text = re.sub(r"\n{3,}", "\n\n", text)   # collapse excess blank lines
    text = text.replace("\f", "\n")            # remove form-feeds
    text = "\n".join(l.rstrip() for l in text.splitlines())
    return text.strip()
